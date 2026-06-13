from __future__ import annotations

import os
import re
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.bullet_sanitize import sanitize_bullet_line
from app.word_swap import enforce_word_swap_budget, is_valid_word_swap


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def plain_text_for_docx(text: str) -> str:
    """
    Strip markdown so Word keeps its original bullets, underlines, and styles.
    Never write *, #, or - bullet prefixes into the .docx.
    """
    if not text:
        return ""

    lines: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        line = re.sub(r"\*([^*]+)\*", r"\1", line)
        line = re.sub(r"__([^_]+)__", r"\1", line)
        line = re.sub(r"_([^_]+)_", r"\1", line)
        line = line.replace("*", "").replace("_", "")
        line = re.sub(r" +", " ", line)
        line = re.sub(r"^[\-\*•]\s+", "", line)
        line = re.sub(r"^\d+\.\s+", "", line)
        line = line.strip()
        if line:
            lines.append(line)

    if not lines:
        return ""
    if len(lines) == 1:
        return lines[0]
    return "\n".join(lines)


def collect_paragraph_slots(doc: Document) -> list[Paragraph]:
    slots: list[Paragraph] = []
    for para in doc.paragraphs:
        slots.append(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    slots.append(para)
    return slots


def get_paragraph_texts(docx_path) -> list[str]:
    doc = Document(docx_path)
    return [p.text.strip() for p in collect_paragraph_slots(doc) if p.text.strip()]


def build_paragraph_prompt(docx_path) -> tuple[str, int]:
    doc = Document(docx_path)
    slots = collect_paragraph_slots(doc)
    lines: list[str] = []
    for i, para in enumerate(slots):
        text = para.text.replace("\n", " ").strip()
        lines.append(f"[P{i}] {text}")
    return "\n".join(lines), len(slots)


def set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace text but keep paragraph style and first run formatting (underline, bold, etc.)."""
    new_text = plain_text_for_docx(new_text)
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_in_paragraph(paragraph: Paragraph, original: str, tailored: str) -> bool:
    """
    Replace original substring with tailored text inside a paragraph.
    Prefers editing a single run to preserve Word formatting.
    """
    original = plain_text_for_docx(original)
    tailored = plain_text_for_docx(tailored)
    if not original or not tailored or original == tailored:
        return False

    full = paragraph.text
    if original in full:
        if full.strip() == original.strip():
            set_paragraph_text(paragraph, tailored)
            return True
        for run in paragraph.runs:
            if original in run.text:
                run.text = run.text.replace(original, tailored, 1)
                return True
        set_paragraph_text(paragraph, full.replace(original, tailored, 1))
        return True

    # Whitespace-tolerant match (extracted text vs Word paragraph spacing)
    if normalize_ws(original) == normalize_ws(full):
        set_paragraph_text(paragraph, tailored)
        return True

    return False


def match_paragraph_text(original: str, slot_texts: list[str]) -> str | None:
    """Find the exact Word paragraph text for a replacement original string."""
    cleaned = plain_text_for_docx(original)
    if not cleaned:
        return None

    norm_orig = normalize_ws(cleaned)
    for slot in slot_texts:
        if slot == cleaned or normalize_ws(slot) == norm_orig:
            return slot

    best: str | None = None
    best_score = 0.0
    for slot in slot_texts:
        score = SequenceMatcher(None, normalize_ws(slot).lower(), norm_orig.lower()).ratio()
        if score > best_score:
            best_score = score
            best = slot
    if best_score >= 0.94:
        return best
    return None


def validate_replacements(
    source_docx_path,
    text_replacements: list[dict],
) -> list[dict]:
    """
    Keep only small word swaps whose original text exists in the uploaded Word file.
    Uses the exact paragraph wording from the .docx as the replacement key.
    """
    slot_texts = get_paragraph_texts(source_docx_path)
    slot_set = set(slot_texts)
    candidates: list[dict] = []
    verb_index = 0

    for item in text_replacements or []:
        if not isinstance(item, dict):
            continue
        tailored = plain_text_for_docx(str(item.get("tailored", "")))
        original = plain_text_for_docx(str(item.get("original", "")))
        if not original or not tailored or original == tailored:
            continue

        slot_match = match_paragraph_text(original, slot_texts)
        if not slot_match or slot_match not in slot_set:
            continue

        tailored, verb_index = sanitize_bullet_line(tailored, verb_index)
        if not is_valid_word_swap(slot_match, tailored):
            continue

        candidates.append({"original": slot_match, "tailored": tailored})

    budgeted = enforce_word_swap_budget(candidates)
    return sorted(budgeted, key=lambda r: len(r["original"]), reverse=True)


def apply_text_replacements(docx_path, replacements: list[dict], out_path=None) -> int:
    path = out_path or docx_path
    doc = Document(docx_path)
    slots = collect_paragraph_slots(doc)
    applied = 0

    for item in replacements:
        original = plain_text_for_docx(str(item.get("original", "")))
        tailored = plain_text_for_docx(str(item.get("tailored", "")))
        if not original or not tailored or original == tailored:
            continue
        for para in slots:
            if replace_in_paragraph(para, original, tailored):
                applied += 1
                break

    doc.save(path)
    return applied


def rebuild_resume_preview(original_resume_text: str, replacements: list[dict]) -> str:
    """Build preview text in the same order as the upload — only reworded lines change."""
    if not replacements:
        return original_resume_text

    lines = original_resume_text.splitlines()
    out_lines: list[str] = []
    replaced: set[str] = set()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            out_lines.append(line)
            continue

        new_line = line
        for item in replacements:
            orig = item["original"]
            if orig in replaced:
                continue
            if stripped == orig or normalize_ws(stripped) == normalize_ws(orig):
                new_line = line.replace(stripped, item["tailored"], 1)
                replaced.add(orig)
                break
        out_lines.append(new_line)

    return "\n".join(out_lines)


def build_tailored_docx(
    source_docx_path,
    tailored_md: str,
    paragraph_updates: list[dict],
    text_replacements: list[dict],
    out_path,
    original_resume_text: str = "",
) -> tuple[int, list[dict], str]:
    """
    Copy the original .docx and apply in-place text replacements only.
    Preserves layout, fonts, bullets, and paragraph order from the upload.
    """
    import shutil

    shutil.copy2(source_docx_path, out_path)

    validated = validate_replacements(source_docx_path, text_replacements or [])
    applied = 0
    if validated:
        applied = apply_text_replacements(out_path, validated, out_path)

    preview = rebuild_resume_preview(original_resume_text, validated)
    ensure_docx_editable(out_path)
    return applied, validated, preview


def ensure_docx_editable(docx_path) -> None:
    """Remove write/edit protection so the downloaded file opens editable in Word."""
    path = Path(docx_path)
    doc = Document(path)
    settings_el = doc.settings.element
    for local_name in (
        "writeProtection",
        "documentProtection",
        "readOnlyRecommended",
        "edit",
    ):
        tag = qn(f"w:{local_name}")
        for child in list(settings_el.findall(tag)):
            settings_el.remove(child)
    doc.save(path)
    os.chmod(path, 0o644)
