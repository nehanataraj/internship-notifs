from io import BytesIO
from pathlib import Path


def extract_text(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in {".txt", ".md", ".markdown"}:
        return content.decode("utf-8", errors="replace").strip()
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext in {".docx", ".doc"}:
        return _extract_docx(content)
    raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    result = "\n\n".join(parts).strip()
    if not result:
        raise ValueError("Could not extract text from PDF (it may be scanned/image-only).")
    return result


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    result = "\n".join(parts).strip()
    if not result:
        raise ValueError("Could not extract text from Word document.")
    return result
